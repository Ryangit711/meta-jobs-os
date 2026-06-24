#!/usr/bin/env python3
"""Generate a universal Master Resume DOCX + PDF — ready to send anytime."""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

ONEDRIVE = "/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0"
LINUX = "/home/aryan/opencode_test/ABHIMANYU-2.0"
FONT = "Calibri"
SIZE = Pt(11)

NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"

def set_margins(doc, margin):
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

def add_hyperlink(p, label, url, font_name, font_size):
    part = p.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    hl_xml = (
        f'<w:hyperlink xmlns:w="{ns_w}" xmlns:r="{ns_r}" '
        f'r:id="{r_id}" w:history="1">'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'<w:color w:val="0563C1"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{label}</w:t>'
        f'</w:r></w:hyperlink>'
    )
    p._p.append(parse_xml(hl_xml))

def add_plain_run(p, text, font_name, font_size, color=None):
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    color_attr = f'<w:color w:val="{color}"/>' if color else ''
    run_xml = (
        f'<w:r xmlns:w="{ns_w}">'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'{color_attr}'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    p._p.append(parse_xml(run_xml))

def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.font.name = FONT
    run.font.size = Pt(12)
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_body(doc, text, bold=False, italic=False, size=None, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = size or SIZE
    run.bold = bold
    run.italic = italic

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        run_b = p.add_run(f"• {bold_prefix}")
        run_b.font.name = FONT
        run_b.font.size = SIZE
        run_b.bold = True
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = SIZE
    else:
        run = p.add_run(f"• {text}")
        run.font.name = FONT
        run.font.size = SIZE

doc = Document()
set_margins(doc, Inches(0.75))

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(NAME)
run.font.name = FONT
run.font.size = Pt(16)
run.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_plain_run(p2, f"{PHONE}  |  ", FONT, Pt(9), color="505050")
add_hyperlink(p2, EMAIL, f"mailto:{EMAIL}", FONT, Pt(9))
add_plain_run(p2, "  |  ", FONT, Pt(9), color="505050")
add_hyperlink(p2, "LinkedIn", f"https://{LINKEDIN}", FONT, Pt(9))
add_plain_run(p2, f"  |  {LOCATION}", FONT, Pt(9), color="505050")

# Summary
add_section_header(doc, "Professional Summary")
add_body(doc,
    "Operations executive who built a multi-site organization from 3 to 70 people across 32 locations, "
    "directed a $17M acquisition exit, and designed the complete operational infrastructure from scratch. "
    "Combines strategic thinking with hands-on execution — equally comfortable leading board-level strategy "
    "sessions, building financial models, managing full P&L ownership, or aligning diverse teams around a "
    "shared plan. Known for creating systems that scale without adding complexity.")

# Core Competencies
add_section_header(doc, "Core Competencies")
add_body(doc,
    "Multi-Site Operations  |  Strategic Planning & Execution  |  P&L Management  |  M&A & Integration  |  "
    "Systems Architecture & Automation  |  Cross-Functional Leadership  |  Board-Level Communication  |  "
    "Quality Improvement & Process Optimization  |  Organizational Design", size=Pt(9.5))

# Experience
add_section_header(doc, "Professional Experience")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(4)
run = p.add_run("SkyflyMD")
run.font.name = FONT
run.font.size = SIZE
run.bold = True
run2 = p.add_run("  |  Director of Operations  |  Phoenix, AZ / Vancouver, BC  |  2017 – 2025")
run2.font.name = FONT
run2.font.size = SIZE

add_body(doc,
    "Led operations for a multi-site organization. Built the operational infrastructure, strategic systems, "
    "and financial frameworks from zero.", italic=True, size=Pt(9.5))

add_bullet(doc,
    " — built a multi-site organization from 3 to 70 employees across 32 locations, designing hiring "
    "frameworks, training programs, operational systems, and quality standards that scaled without adding complexity",
    bold_prefix="Organizational Leadership")
add_bullet(doc,
    " — managed full P&L ownership for $4M organization: budget planning, variance analysis, resource "
    "allocation, capital expenditure planning across 12 departments",
    bold_prefix="Financial Management")
add_bullet(doc,
    " — designed and implemented complete operational tech stack (EHR, billing, scheduling, analytics, "
    "reporting) — automated workflows, KPI dashboards, real-time performance visibility across all locations",
    bold_prefix="Systems Architecture")
add_bullet(doc,
    " — directed end-to-end $17M acquisition: structured 8 due diligence workstreams, managed integration "
    "across all locations, consolidated 8 operational systems within 90 days, retained 100% of key talent",
    bold_prefix="M&A & Integration")
add_bullet(doc,
    " — built governance rhythms, communication frameworks, and decision-making structures that aligned "
    "diverse stakeholders across multiple locations, functional areas, and external partners",
    bold_prefix="Cross-Functional Leadership")

# Earlier
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Earlier Career")
run.font.name = FONT
run.font.size = SIZE
run.bold = True
add_bullet(doc, "Digital Strategy Manager (2016–2018) — digital strategy, campaign analytics, ROI measurement")
add_bullet(doc, "Client Services Representative (2014–2016) — client escalations, complex issue resolution")

# Education
add_section_header(doc, "Education")
add_body(doc, "Master of Business Administration (MBA)", bold=True, space_after=0)
add_body(doc, "Post-Baccalaureate Diploma in Business Management — KPU, Surrey, BC", size=Pt(9), space_after=0)
add_body(doc, "Bachelor of Science, Information Technology", size=Pt(9), space_after=0)

# Technical
add_section_header(doc, "Technical Proficiency")
add_body(doc,
    "Financial Modeling & Analysis  |  OKR Frameworks  |  KPI Dashboards & Data Visualization  |  "
    "Google Workspace  |  CRM Platforms  |  Project Management Tools  |  AI-Augmented Workflows", size=Pt(9))

# Save
folder = f"{ONEDRIVE}/Master"
lfolder = f"{LINUX}/Master"
os.makedirs(folder, exist_ok=True)
os.makedirs(lfolder, exist_ok=True)

docx_path = os.path.join(folder, "Aman_Kumar_Master_Resume.docx")
doc.save(docx_path)
doc.save(os.path.join(lfolder, "Aman_Kumar_Master_Resume.docx"))
print(f"DOCX: {docx_path}")

# Generate PDF via WeasyPrint
html_parts = [f'''<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
@page {{ margin: 0.75in; size: letter; }}
body {{ font-family: Calibri, sans-serif; font-size: 11pt; color: #000; line-height: 1.15; }}
h1 {{ font-size: 16pt; font-weight: bold; text-align: center; margin: 0; padding: 0; }}
.contact {{ text-align: center; font-size: 9pt; color: #505050; margin: 0 0 4pt 0; }}
.section-header {{ font-size: 12pt; font-weight: bold; border-bottom: 1px solid #000; margin-top: 10pt; margin-bottom: 3pt; text-transform: uppercase; }}
.body-text {{ font-size: 11pt; margin: 0 0 2pt 0; }}
.bullet {{ font-size: 11pt; margin: 0 0 1pt 0; padding-left: 25pt; text-indent: -25pt; }}
.bold {{ font-weight: bold; }}
.italic {{ font-style: italic; }}
</style></head><body>
<h1>Aman Kumar</h1>
<p class="contact">+1 236-885-2285  |  amankumar7111@outlook.com  |  linkedin.com/in/aman1776  |  Vancouver, BC</p>

<p class="section-header">Professional Summary</p>
<p class="body-text">Operations executive who built a multi-site organization from 3 to 70 people across 32 locations, directed a $17M acquisition exit, and designed the complete operational infrastructure from scratch. Combines strategic thinking with hands-on execution — equally comfortable leading board-level strategy sessions, building financial models, managing full P&amp;L ownership, or aligning diverse teams around a shared plan. Known for creating systems that scale without adding complexity.</p>

<p class="section-header">Core Competencies</p>
<p class="body-text">Multi-Site Operations  |  Strategic Planning &amp; Execution  |  P&amp;L Management  |  M&amp;A &amp; Integration  |  Systems Architecture &amp; Automation  |  Cross-Functional Leadership  |  Board-Level Communication  |  Quality Improvement &amp; Process Optimization  |  Organizational Design</p>

<p class="section-header">Professional Experience</p>
<p class="body-text"><b>SkyflyMD</b>  |  Director of Operations  |  Phoenix, AZ / Vancouver, BC  |  2017 – 2025</p>
<p class="body-text"><i>Led operations for a multi-site organization. Built the operational infrastructure, strategic systems, and financial frameworks from zero.</i></p>
<p class="bullet"><span class="bold">Organizational Leadership</span> — built a multi-site organization from 3 to 70 employees across 32 locations, designing hiring frameworks, training programs, operational systems, and quality standards that scaled without adding complexity</p>
<p class="bullet"><span class="bold">Financial Management</span> — managed full P&amp;L ownership for $4M organization: budget planning, variance analysis, resource allocation, capital expenditure planning across 12 departments</p>
<p class="bullet"><span class="bold">Systems Architecture</span> — designed and implemented complete operational tech stack (EHR, billing, scheduling, analytics, reporting) — automated workflows, KPI dashboards, real-time performance visibility across all locations</p>
<p class="bullet"><span class="bold">M&amp;A &amp; Integration</span> — directed end-to-end $17M acquisition: structured 8 due diligence workstreams, managed integration across all locations, consolidated 8 operational systems within 90 days, retained 100% of key talent</p>
<p class="bullet"><span class="bold">Cross-Functional Leadership</span> — built governance rhythms, communication frameworks, and decision-making structures that aligned diverse stakeholders across multiple locations, functional areas, and external partners</p>

<p class="body-text"><b>Earlier Career</b></p>
<p class="bullet">Digital Strategy Manager (2016–2018) — digital strategy, campaign analytics, ROI measurement</p>
<p class="bullet">Client Services Representative (2014–2016) — client escalations, complex issue resolution</p>

<p class="section-header">Education</p>
<p class="body-text"><b>Master of Business Administration (MBA)</b></p>
<p class="body-text" style="font-size: 9pt;">Post-Baccalaureate Diploma in Business Management — KPU, Surrey, BC</p>
<p class="body-text" style="font-size: 9pt;">Bachelor of Science, Information Technology</p>

<p class="section-header">Technical Proficiency</p>
<p class="body-text" style="font-size: 9pt;">Financial Modeling &amp; Analysis  |  OKR Frameworks  |  KPI Dashboards &amp; Data Visualization  |  Google Workspace  |  CRM Platforms  |  Project Management Tools  |  AI-Augmented Workflows</p>

</body></html>''']

html_str = '\n'.join(html_parts)
pdf_path = os.path.join(folder, "Aman_Kumar_Master_Resume.pdf")
from weasyprint import HTML
HTML(string=html_str).write_pdf(pdf_path)
from shutil import copy2
copy2(pdf_path, os.path.join(lfolder, "Aman_Kumar_Master_Resume.pdf"))
print(f"PDF:  {pdf_path}")
