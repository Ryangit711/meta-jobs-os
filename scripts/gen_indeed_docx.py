#!/usr/bin/env python3
"""Generate Indeed resume + cover letter DOCX files - ready to submit."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUT_DIR = "/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0/2026-06-21/Indeed"
LINUX_DIR = "/home/aryan/opencode_test/ABHIMANYU-2.0/2026-06-21/Indeed"
os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(LINUX_DIR, exist_ok=True)

FONT_NAME = "Liberation Sans"
FONT_SIZE = Pt(10)
HEADER_SIZE = Pt(12)
TITLE_SIZE = Pt(14)
MARGINS = Inches(0.75)
NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"


def set_margins(doc, margin):
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def add_contact(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{NAME}")
    run.font.name = FONT_NAME
    run.font.size = TITLE_SIZE
    run.bold = True
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{PHONE} | {EMAIL} | {LINKEDIN} | {LOCATION}")
    run2.font.name = FONT_NAME
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(80, 80, 80)
    
    return p2


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = FONT_NAME
    run.font.size = HEADER_SIZE
    run.bold = True
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_body_text(doc, text, bold=False, italic=False, space_after=2, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        run_b = p.add_run(f"• {bold_prefix}")
        run_b.font.name = FONT_NAME
        run_b.font.size = FONT_SIZE
        run_b.bold = True
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
    else:
        run = p.add_run(f"• {text}")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
    return p


def generate_resume():
    doc = Document()
    set_margins(doc, MARGINS)
    
    add_contact(doc)
    
    # Professional Summary
    add_section_header(doc, "Professional Summary")
    add_body_text(doc, 
        "Operations executive who built multi-site operational infrastructure from scratch, "
        "scaling a healthcare organization from 3 to 70 employees across 32 locations, then "
        "directed the full-cycle execution of a $17M acquisition. Combines strategic thinking "
        "with hands-on execution — equally comfortable building financial models, designing "
        "governance rhythms, and leading cross-functional teams through complex integration programs."
    )
    
    # Core Competencies
    add_section_header(doc, "Core Competencies")
    add_body_text(doc,
        "M&A Integration & Execution  |  Cross-Functional Program Management  |  "
        "Operational Infrastructure Design  |  Strategic Planning & OKRs  |  "
        "P&L Management & Financial Modeling  |  Board & Executive Reporting  |  "
        "Multi-Site Operations  |  Change Management  |  Due Diligence  |  Playbook Creation",
        size=Pt(9)
    )
    
    # Professional Experience
    add_section_header(doc, "Professional Experience")
    
    # SkyflyMD
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("SkyflyMD")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = True
    run2 = p.add_run("  |  Director of Operations  |  Phoenix, AZ / Vancouver, BC  |  2017 – 2025")
    run2.font.name = FONT_NAME
    run2.font.size = FONT_SIZE
    
    add_body_text(doc, 
        "Directed end-to-end operations for a multi-site healthcare group. Served as the primary "
        "bridge between executive leadership and all operational teams — translating strategic "
        "direction into executable plans across clinical operations, scheduling, finance, quality "
        "assurance, and provider relations. Built every system from zero — no playbook existed.",
        italic=True, size=Pt(9)
    )
    
    add_bullet(doc, 
        " — structured 8 concurrent due diligence workstreams spanning finance, legal, operations, "
        "and provider contracts. Built end-to-end integration playbook with Day 1 and Day 100 "
        "milestones. Retained 100% of key talent through the $17M transition.",
        bold_prefix="Full-Cycle Acquisition Execution"
    )
    add_bullet(doc,
        " — designed communication protocols, decision-making frameworks, leadership meeting "
        "cadences, escalation paths, and department playbooks across 5 clinic groups in multiple states.",
        bold_prefix="Multi-Site Operating System"
    )
    add_bullet(doc,
        " — built hiring frameworks, role families, training programs with sub-30-day ramp time, "
        "cross-border team coordination (India → US → Canada), and 100% quality standards.",
        bold_prefix="Organizational Infrastructure"
    )
    add_bullet(doc,
        " — designed and managed company-wide OKR system across 5 consecutive planning cycles "
        "with consistent attainment across all operational metrics. Built board-level reporting "
        "templates, financial models, and departmental budget ownership across 12 departments.",
        bold_prefix="Strategic Planning & OKR System"
    )
    add_bullet(doc,
        " — led technology transformation across all 32 locations: EHR migration, RCM system "
        "implementation, billing automation, scheduling optimization, and telemedicine deployment.",
        bold_prefix="Technology Transformation"
    )
    
    # Earlier Experience
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Earlier Career")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = True
    
    add_bullet(doc,
        " — led digital strategy, campaign analytics, and ROI measurement for multi-channel marketing programs (2016–2018).",
        bold_prefix="Digital Marketing Manager"
    )
    add_bullet(doc,
        " — managed client escalations, resolved complex service issues, and maintained satisfaction metrics (2014–2016).",
        bold_prefix="Client Services Representative"
    )
    
    # Education
    add_section_header(doc, "Education")
    add_body_text(doc, 
        "Master of Business Administration (MBA)", bold=True, space_after=0)
    add_body_text(doc,
        "Post-Baccalaureate Diploma in Business Management — Kwantlen Polytechnic University, Surrey, BC",
        size=Pt(9), space_after=0)
    add_body_text(doc,
        "Bachelor of Science, Information Technology",
        size=Pt(9), space_after=0)
    
    # Technical Proficiency
    add_section_header(doc, "Technical Proficiency")
    add_body_text(doc,
        "Athenahealth  |  eClinicalWorks  |  CRM Platforms  |  Google Workspace  |  "
        "Financial Modeling & Analysis  |  OKR Frameworks  |  Jira / Confluence  |  Data Visualization",
        size=Pt(9)
    )
    
    path = os.path.join(OUT_DIR, "Aman_Kumar_Indeed_SrMgr_Integration.docx")
    doc.save(path)
    # Also copy to Linux dir
    doc.save(os.path.join(LINUX_DIR, "Aman_Kumar_Indeed_SrMgr_Integration.docx"))
    print(f"Resume saved to {path}")
    return path


def generate_cover_letter():
    doc = Document()
    set_margins(doc, MARGINS)
    
    # Header
    add_body_text(doc, NAME, bold=True, space_after=0)
    add_body_text(doc, f"{PHONE} | {EMAIL} | {LINKEDIN}", size=Pt(9), space_after=0)
    add_body_text(doc, LOCATION, size=Pt(9), space_after=4)
    add_body_text(doc, "June 21, 2026", space_after=8)
    
    add_body_text(doc, "Indeed", bold=False, space_after=0)
    add_body_text(doc, "Vancouver, BC", space_after=8)
    
    add_body_text(doc, "Re: Sr. Manager, Integration & Business Acceleration — Reference ID: 47053", 
                  bold=True, space_after=8)
    
    add_body_text(doc, "Dear Hiring Manager,", space_after=4)
    
    body = (
        "Your mission is to help people get jobs. Mine is to build the operational infrastructure "
        "that makes organizations scalable, efficient, and ready for their next phase of growth. "
        "When I read the Sr. Manager, Integration & Business Acceleration role, I recognized a "
        "team that needs exactly what I have spent eight years doing — building from zero, leading "
        "complex cross-functional programs, and ensuring that strategic moves deliver their full value.\n\n"
        
        "I joined SkyflyMD when it was 3 people in a single location. When I left, it was 70 people "
        "across 32 locations, operating with systems I designed from scratch. The defining moment of "
        "that journey was directing our $17M acquisition — structuring 8 concurrent due diligence "
        "workstreams across finance, legal, operations, and provider contracts. I built the integration "
        "playbook, tracked Day 1 readiness and Day 100 milestones, consolidated 8 separate operational "
        "systems into one unified platform, and retained 100% of our key talent through the transition.\n\n"
        
        "What made that possible was not a pre-existing framework. There was no playbook. I created it — "
        "the governance rhythms, the reporting cadences, the escalation paths, the decision-making "
        "frameworks — all from scratch, across 5 clinic groups operating under different state regulations. "
        "I learned that successful integration depends not on the complexity of the plan, but on creating "
        "simple, clear structures that diverse teams can rally around, combined with the discipline to track "
        "progress until synergy targets are met.\n\n"
        
        "What drew me to Indeed is that you are data-driven without being jargon-heavy, and mission-focused "
        "without being sentimental. That is how I operate. I believe in measuring what matters, building "
        "systems that outlast the people who create them, and keeping things simple enough that they actually "
        "get used. Your 'job seeker first' value resonates because it mirrors how I have always made "
        "operational decisions — starting with the end-user and working backward.\n\n"
        
        "I am not looking for a role that requires a playbook to exist before I start. I am looking for "
        "one that needs a playbook written. The Integration & Business Acceleration team at Indeed is "
        "exactly that kind of environment.\n\n"
        
        "I would welcome the opportunity to discuss how my experience building and integrating multi-site "
        "operations can support Indeed's continued growth through M&A. Thank you for your time and consideration."
    )
    
    add_body_text(doc, body, space_after=8)
    add_body_text(doc, "Best regards,", space_after=2)
    add_body_text(doc, NAME, bold=True, space_after=0)
    add_body_text(doc, PHONE, size=Pt(9), space_after=0)
    add_body_text(doc, EMAIL, size=Pt(9))
    
    path = os.path.join(OUT_DIR, "Cover_Letter_Indeed_SrMgr_Integration.docx")
    doc.save(path)
    doc.save(os.path.join(LINUX_DIR, "Cover_Letter_Indeed_SrMgr_Integration.docx"))
    print(f"Cover letter saved to {path}")
    return path


if __name__ == "__main__":
    generate_resume()
    generate_cover_letter()
    print("Done.")
