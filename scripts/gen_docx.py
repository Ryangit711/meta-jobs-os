#!/usr/bin/env python3
"""Parametric DOCX generator — reads company fit map + SHOOT package, outputs tailored DOCX."""

import sys
import os
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree

ONEDRIVE = "/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0"
LINUX = "/home/aryan/opencode_test/ABHIMANYU-2.0"

NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"

# Company-specific configs
CONFIG = {
    "Indeed": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 1,
        "ats_notes": "Indeed internal ATS — DOCX preferred, Liberation Sans, 0.75in margins"
    },
    "Methanex": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Methanex career portal — DOCX, Calibri 11pt, 2-page for Director level"
    },
    "Deloitte": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Deloitte Workday ATS — DOCX, Calibri 10pt"
    },
    "Hiive": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Hiive Ashby ATS — DOCX, Calibri 11pt"
    },
    "Providence_Healthcare": {
        "font": "Calibri",
        "size": Pt(11),
        "header_size": Pt(13),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Providence Oracle Cloud ATS — DOCX, Calibri 11pt"
    },
    "DoorDash_Canada": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 1,
        "ats_notes": "DoorDash Greenhouse ATS — DOCX, Liberation Sans 10pt, 0.75in margins"
    },
    "UBC": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 1,
        "ats_notes": "UBC / Indeed ATS — DOCX, Liberation Sans 10pt, 0.75in margins, public sector format"
    },
    "Practice_Better": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "Practice Better Greenhouse ATS — DOCX, Liberation Sans 10pt, 0.75in margins"
    },
    "RAM_Consulting": {
        "font": "Liberation Sans",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "RAM Consulting DayforceHCM ATS — DOCX, Liberation Sans 10pt, 0.75in margins, project delivery focus"
    },
    "BWZ": {
        "font": "Calibri",
        "size": Pt(10),
        "header_size": Pt(12),
        "margins": Inches(0.75),
        "pages": 2,
        "ats_notes": "BWZ Lever ATS — DOCX, Calibri 10pt, 0.75in margins, strategy & operations focus"
    }
}

def set_margins(doc, margin):
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

def add_contact(doc, config):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(NAME)
    run.font.name = config["font"]
    run.font.size = Pt(16)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_plain_run(p2, f"{PHONE}  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(p2, EMAIL, f"mailto:{EMAIL}", config["font"], Pt(9))
    add_plain_run(p2, "  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(p2, "LinkedIn", f"https://{LINKEDIN}", config["font"], Pt(9))
    add_plain_run(p2, f"  |  {LOCATION}", config["font"], Pt(9), color="505050")

def add_section_header(doc, text, config):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.font.name = config["font"]
    run.font.size = config["header_size"]
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_body(doc, text, config, bold=False, italic=False, size=None, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = config["font"]
    run.font.size = size or config["size"]
    run.bold = bold
    run.italic = italic

def add_bullet(doc, text, config, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    if bold_prefix:
        run_b = p.add_run(f"• {bold_prefix}")
        run_b.font.name = config["font"]
        run_b.font.size = config["size"]
        run_b.bold = True
        run = p.add_run(text)
        run.font.name = config["font"]
        run.font.size = config["size"]
    else:
        run = p.add_run(f"• {text}")
        run.font.name = config["font"]
        run.font.size = config["size"]

def add_hyperlink_contact(p, label, url, font_name, font_size):
    part = p.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    hyperlink_xml = (
        f'<w:hyperlink xmlns:w="{ns_w}" xmlns:r="{ns_r}" '
        f'r:id="{r_id}" w:history="1">'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'<w:color w:val="0563C1"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{label}</w:t>'
        f'</w:r></w:hyperlink>'
    )
    hyperlink_elem = parse_xml(hyperlink_xml)
    p._p.append(hyperlink_elem)

def add_plain_run(p, text, font_name, font_size, color=None):
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    color_attr = f'<w:color w:val="{color}"/>' if color else ''
    run_xml = (
        f'<w:r xmlns:w="{ns_w}">'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
        f'<w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'{color_attr}'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    run_elem = parse_xml(run_xml)
    p._p.append(run_elem)

def add_signature(doc, config):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    add_plain_run(p, "Best regards,", config["font"], config["size"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    add_plain_run(p2, NAME, config["font"], config["size"])
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)
    add_plain_run(p3, PHONE, config["font"], config["size"], color="505050")
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(0)
    p4.paragraph_format.space_before = Pt(0)
    add_hyperlink_contact(p4, EMAIL, f"mailto:{EMAIL}", config["font"], config["size"])

def get_date(company):
    DATES = {
        "Indeed": "2026-06-20",
        "Methanex": "2026-06-21",
        "Deloitte": "2026-06-19",
        "Hiive": "2026-06-22",
        "Providence_Healthcare": "2026-06-22",
        "DoorDash_Canada": "2026-06-24",
        "UBC": "2026-06-25",
        "Practice_Better": "2026-06-25",
        "RAM_Consulting": "2026-06-25",
        "BWZ": "2026-06-25"
    }
    return DATES.get(company, "2026-06-22")

def generate(company):
    config = CONFIG.get(company, CONFIG["Methanex"])
    date_str = get_date(company)
    folder = f"{ONEDRIVE}/{date_str}/{company}"
    lfolder = f"{LINUX}/{date_str}/{company}"
    os.makedirs(folder, exist_ok=True)
    os.makedirs(lfolder, exist_ok=True)

    # --- RESUME ---
    doc = Document()
    set_margins(doc, config["margins"])
    add_contact(doc, config)

    add_section_header(doc, "Professional Summary", config)
    if company == "Methanex":
        add_body(doc,
            "Strategy and operations executive who built the strategic infrastructure for a multi-site "
            "organization from zero, scaling it from 3 to 70 employees across 32 locations and directing "
            "a $17M acquisition. Combines board-level strategic thinking with hands-on financial modelling, "
            "M&A execution, and cross-functional leadership. Equally comfortable leading an ELT strategy "
            "session, building a valuation model, or aligning diverse teams around a shared plan.", config)
    elif company == "Hiive":
        add_body(doc,
            "Operations strategist and systems builder who designed and implemented the complete operational "
            "infrastructure for a startup that grew 23x — from 3 to 70 employees across 32 locations — "
            "and directed its $17M exit. Combines a builder's instinct for scalable systems with an operator's "
            "discipline for data integrity, revenue lifecycle optimization, and cross-functional execution.", config)
    elif company == "Providence_Healthcare":
        add_body(doc,
            "Healthcare operations leader with 8 years of progressive experience directing multi-site clinical "
            "and operational infrastructure across 32 locations and 12 departments. Scaled a healthcare startup "
            "from 3 to 70 employees and $4M ARR, managed full P&L ownership, led quality improvement initiatives, "
            "and directed a $17M acquisition end-to-end. Combines strategic thinking with hands-on operational "
            "execution in complex, multi-stakeholder healthcare environments.", config)
    elif company == "DoorDash_Canada":
        add_body(doc,
            "Operations executive with 8 years scaling a business from 3 to 70 people, 32 locations, and $4M ARR. "
            "Led the complete technology transformation, owned end-to-end P&L, and delivered a $17M exit. "
            "Built marketplace-style incentive systems balancing worker earnings with cost efficiency — directly transferable to Dasher pay strategy. "
            "Also led cross‑functional teams across 12 departments, delivering board‑level insights and strategic alignment.", config)
    elif company == "Practice_Better":
        add_body(doc,
            "Revenue operations executive who architected the complete RevOps infrastructure for a multi-site healthcare "
            "SaaS business from zero to $4M ARR across 32 locations, then directed through a $17M exit. "
            "Combines a builder's instinct for scalable systems with an operator's discipline for SaaS metrics (MRR, ARR, churn, CAC, LTV), "
            "pipeline integrity, and cross-functional GTM execution. Proven ability to own the revenue operations function end-to-end "
            "for a scaling healthcare SaaS company.", config)
    elif company == "BWZ":
        add_body(doc,
            "Operations executive who built a multi-site organization from 3 to 70 people, 32 locations, and $4M ARR — "
            "then directed a $17M exit. Specializes in zero-to-one initiative build, cross-functional program discipline, "
            "and business case development for scale-stage companies. Combines a builder's hands-on execution with "
            "a strategist's ability to pressure-test ideas and turn ambiguity into actionable plans. "
            "MBA preferred, founder DNA, AI-native operator.", config)
    else:
        add_body(doc,
            "Operations executive who built multi-site operational infrastructure from scratch, "
            "scaling from 3 to 70 employees across 32 locations, then directed a $17M acquisition. "
            "Combines strategic thinking with execution.", config)

    add_section_header(doc, "Core Competencies", config)
    if company == "Methanex":
        add_body(doc,
            "Corporate Strategy & Planning  |  Financial Modelling & Valuation  |  M&A Execution & Integration  |  "
            "Board & Executive Communication  |  Cross-Functional Leadership  |  OKR & Performance Systems  |  "
            "Capital Allocation  |  Strategic Growth Initiatives", config, size=Pt(9.5))
    elif company == "Hiive":
        add_body(doc,
            "Revenue Operations  |  Operational Infrastructure Design  |  Systems Architecture & Automation  |  "
            "Bottleneck Analysis  |  Workflow Optimization  |  Data Integrity & Reporting  |  "
            "Cross-Functional Execution  |  M&A & Strategic Projects", config, size=Pt(9.5))
    elif company == "Providence_Healthcare":
        add_body(doc,
            "Multi-Site Healthcare Operations  |  Clinical Operations Leadership  |  Quality Improvement  |  "
            "Financial Management & Budgeting  |  Interdisciplinary Team Leadership  |  "
            "Change Management & Transformation  |  Regulatory Compliance  |  Strategic Planning & OKRs", config, size=Pt(9.5))
    elif company == "DoorDash_Canada":
        add_body(doc,
            "Marketplace Operations  |  Pay & Incentive Design  |  P&L Management  |  "
            "Cross-Functional Leadership  |  Behavioral Economics  |  Data-Driven Strategy  |  "
            "0-to-1 Scaling  |  Operational Excellence", config, size=Pt(9.5))
    elif company == "Practice_Better":
        add_body(doc,
            "Revenue Operations  |  SaaS Metrics (MRR, ARR, Churn, CAC, LTV)  |  P&L Management  |  "
            "Revenue Cycle Management  |  GTM Alignment  |  Pipeline Management  |  Forecasting & Budgeting  |  "
            "Process Optimization  |  Workflow Automation  |  EHR / Practice Management Platforms  |  "
            "KPI Dashboard Design  |  SaaS Scaling Infrastructure  |  Team Building (3$\\rightarrow$70 FTEs)", config, size=Pt(9.5))
    elif company == "BWZ":
        add_body(doc,
            "Zero-to-One Initiative Build  |  Program & Project Discipline  |  Business Case Development  |  "
            "Cross-Functional Leadership  |  Business Modeling (Revenue, Cost, Margin)  |  P&L Management  |  "
            "AI-Augmented Workflows  |  Data Analytics & Dashboarding  |  Strategic Planning & OKRs  |  "
            "Operational Infrastructure Design", config, size=Pt(9.5))
    else:
        add_body(doc,
            "M&A Integration  |  Cross-Functional Program Management  |  Operational Infrastructure  |  "
            "Strategic Planning & OKRs  |  P&L Management  |  Board Reporting  |  Multi-Site Operations", config, size=Pt(9.5))

    add_section_header(doc, "Professional Experience", config)

    # SkyflyMD
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("SkyflyMD")
    run.font.name = config["font"]
    run.font.size = config["size"]
    run.bold = True
    run2 = p.add_run("  |  Director of Operations  |  Phoenix, AZ / Vancouver, BC  |  2017 – 2025")
    run2.font.name = config["font"]
    run2.font.size = config["size"]

    if company == "Methanex":
        add_body(doc,
            "Led strategy and operations for a multi-site healthcare group. Built the strategic planning "
            "process, financial infrastructure, and operational systems from zero.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — built company-wide strategic planning framework across 5 annual cycles: board-level strategy "
            "sessions, departmental OKR cascades, quarterly performance reviews with executive leadership",
            config, bold_prefix="Strategic Planning & Execution")
        add_bullet(doc,
            " — multi-scenario P&L models, capital allocation frameworks, departmental budgets across "
            "12 departments, board-ready reporting templates",
            config, bold_prefix="Financial Modelling")
        add_bullet(doc,
            " — directed end-to-end $17M acquisition: 8 due diligence workstreams, Day 1 readiness, "
            "100% key talent retention, 90-day systems consolidation",
            config, bold_prefix="M&A & Integration")
        add_bullet(doc,
            " — organizational infrastructure for 70 employees across 5 clinic groups, 32 locations — "
            "hiring frameworks, training, quality standards, cross-border coordination",
            config, bold_prefix="Organizational Leadership")
    elif company == "Hiive":
        add_body(doc,
            "Built the complete operational infrastructure for a multi-site healthcare startup from zero. "
            "Served as the primary systems architect, process designer, and cross-functional integrator "
            "across all revenue-generating and operational functions.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — from zero: designed and implemented the entire operational tech stack (EHR, billing, scheduling, "
            "RCM, analytics) that scaled from 3 to 70 employees across 32 locations without adding complexity",
            config, bold_prefix="Systems Architecture & Automation")
        add_bullet(doc,
            " — identified and eliminated bottlenecks in the revenue lifecycle — automated billing workflows, "
            "replaced manual reconciliation with real-time RCM, reduced admin overhead by 40%+",
            config, bold_prefix="Revenue Operations & Optimization")
        add_bullet(doc,
            " — built KPI dashboards, board-level reporting, and data integrity systems from scratch — "
            "replaced manual spreadsheets with real-time analytics across all 32 locations",
            config, bold_prefix="Data Infrastructure & Reporting")
        add_bullet(doc,
            " — directed full-cycle $17M acquisition: 8 diligence workstreams, integration playbook, "
            "consolidated 8 systems within 90 days, retained 100% of key talent",
            config, bold_prefix="M&A & Strategic Projects")
    elif company == "Providence_Healthcare":
        add_body(doc,
            "Directed clinical and operational leadership for a multi-site healthcare organization across 32 locations. "
            "Served as the primary operational leader coordinating interdisciplinary teams, managing resources, "
            "and driving quality improvement across 12 departments and 5 clinic groups.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — led clinical and operational leadership across 32 multi-site locations, coordinating interdisciplinary "
            "teams to deliver integrated, patient-centered care across 5 clinic groups and 12 departments",
            config, bold_prefix="Multi-Site Healthcare Operations")
        add_bullet(doc,
            " — managed full P&L ownership for $4M revenue healthcare organization — budget planning, variance analysis, "
            "resource allocation, capital expenditure planning across 32 locations",
            config, bold_prefix="Financial Management & Resource Allocation")
        add_bullet(doc,
            " — drove quality improvement initiatives reducing administrative overhead by 40%+ through workflow "
            "automation, standardized processes, and data-driven performance management",
            config, bold_prefix="Quality Improvement & Process Optimization")
        add_bullet(doc,
            " — directed end-to-end $17M acquisition: 8 diligence workstreams, Day 1 readiness, 90-day systems "
            "consolidation across 32 locations, 100% key talent retention",
            config, bold_prefix="Change Management & Transformation")
        add_bullet(doc,
            " — managed full operational infrastructure for a geriatrics-specialized multi-site practice — scheduling, "
            "billing (ICD coding, insurance claims), multi-facility coordination across 5+ clinic locations and "
            "multiple senior care homes serving aging populations",
            config, bold_prefix="Geriatric Practice & Senior Care Operations")
    elif company == "DoorDash_Canada":
        add_body(doc,
            "Built the business from zero — 3 to 70 people, 32 locations, $4M ARR. Owned end-to-end "
            "P&L, designed the incentive architecture, led the technology transformation, and delivered "
            "a $17M exit.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — scaled 3 to 70 people, 32 locations — built infrastructure, processes, and hiring "
            "frameworks that supported 23x growth without adding complexity",
            config, bold_prefix="Operational Scaling & Infrastructure")
        add_bullet(doc,
            " — designed incentive compensation across 32 markets — tuned per location for labour "
            "cost, competitive pressure, and worker expectations while maintaining budget discipline",
            config, bold_prefix="Pay & Incentive Design")
        add_bullet(doc,
            " — directed full-cycle acquisition: 8 diligence workstreams, Day 1 readiness, 90-day "
            "systems consolidation, 100% key talent retention",
            config, bold_prefix="M&A & Exit Execution")
        add_bullet(doc,
            " — managed $3M+ annual budget, multi-scenario forecasting, variance analysis, capital "
            "allocation across 12 departments — presented board-ready reporting to investors",
            config, bold_prefix="P&L Management & Financial Operations")
        add_bullet(doc,
            " — led transformation from paper to fully integrated digital platform — EHR, billing, "
            "scheduling, analytics — selected, deployed, and owned every system",
            config, bold_prefix="Technology Transformation")
        add_bullet(doc,
            " — implemented KPI dashboards across 32 locations, improving operational visibility and reducing reporting lag by 30%",
            config, bold_prefix="KPI Dashboards")
        add_bullet(doc,
            " — negotiated vendor contracts and technology procurement, achieving 12% cost savings while maintaining service quality",
            config, bold_prefix="Vendor Management & Cost Savings")
    elif company == "UBC":
        add_body(doc,
            "Directed end-to-end operations for a multi-site healthcare organization. Served as the primary "
            "strategic and operational leader, building governance frameworks, stakeholder engagement "
            "processes, and financial infrastructure from zero.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — built company-wide strategic planning framework across 5 annual cycles: board-level strategy "
            "sessions, departmental OKR cascades, quarterly performance reviews — delivering 100% alignment "
            "on annual priorities across 5 clinic groups and 12 departments",
            config, bold_prefix="Strategic Planning & Governance")
        add_bullet(doc,
            " — managed full P&L ownership for $4M ARR organization: department-level budgets, variance "
            "analysis, multi-scenario forecasting, and board-ready financial reporting across 12 departments",
            config, bold_prefix="Budget & Financial Management")
        add_bullet(doc,
            " — directed end-to-end $17M acquisition: 8 concurrent due diligence workstreams, Day 1 readiness, "
            "90-day systems consolidation across 32 locations, 100% key talent retention",
            config, bold_prefix="Acquisition & Integration Leadership")
        add_bullet(doc,
            " — built operational governance framework satisfying both internal accountability requirements and "
            "external regulatory review (HIPAA, state licensing) across 5 jurisdictions",
            config, bold_prefix="Compliance & Regulatory Governance")
        add_bullet(doc,
            " — stakeholder engagement across 12 departments, 5 clinic groups with competing priorities: "
            "established monthly check-in cadences, escalation frameworks, and prioritization matrices",
            config, bold_prefix="Stakeholder Engagement & Alignment")
    elif company == "Practice_Better":
        add_body(doc,
            "Architected the complete revenue operations infrastructure for a multi-site healthcare SaaS business "
            "from zero to \$4M ARR across 32 locations — built the systems, teams, and processes that converted "
            "pipeline into predictable, scalable revenue.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — built end-to-end revenue operations from zero: EHR, billing, scheduling, RCM, and analytics — "
            "designed and implemented the entire tech stack from scratch — scaled from 3 to 70 employees "
            "across 32 locations without adding operational complexity",
            config, bold_prefix="Revenue Operations Infrastructure")
        add_bullet(doc,
            " — architected SaaS metrics framework: MRR, ARR, churn, CAC, LTV, and payback period — "
            "built KPI dashboards replacing manual spreadsheets — reduced reporting lag by 30%, "
            "enabled real-time data-driven decisions across all 32 locations",
            config, bold_prefix="SaaS Metrics & Data Infrastructure")
        add_bullet(doc,
            " — identified and eliminated revenue lifecycle bottlenecks: automated billing workflows, "
            "replaced manual reconciliation with real-time RCM, reduced administrative overhead by 40%+",
            config, bold_prefix="Pipeline Management & Optimization")
        add_bullet(doc,
            " — managed full P&L ownership for \$4M ARR organization: budget planning, variance analysis, "
            "resource allocation, and multi-scenario forecasting across 12 departments",
            config, bold_prefix="P&L Ownership & Forecasting")
        add_bullet(doc,
            " — directed full-cycle \$17M acquisition: 8 due diligence workstreams, integration playbook, "
            "Day 1/100 milestones, systems consolidation across 8$\\rightarrow$1 unified platform, "
            "100% key talent retention",
            config, bold_prefix="M&A Revenue Integration")
        add_bullet(doc,
            " — coordinated cross-functional GTM alignment: sales, operations, finance, and clinical teams — "
            "reduced pipeline variance from 40% to under 10% through unified reporting and accountability",
            config, bold_prefix="GTM Alignment & Cross-Functional Leadership")
    elif company == "BWZ":
        add_body(doc,
            "Built a multi-site organization from 3 to 70 people, 32 locations, and $4M ARR — then directed "
            "a $17M exit. Served as the primary operator, systems architect, and cross-functional integrator "
            "across all revenue-generating and operational functions.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — owned end-to-end build of zero-to-one initiatives: took new ventures and expansions from idea "
            "to launch — costing, lead time, pricing, go-to-market — across 32 locations and 5 clinic groups",
            config, bold_prefix="Zero-to-One Initiative Build")
        add_bullet(doc,
            " — installed project discipline across leadership initiatives: designed governance frameworks that "
            "ensured business cases were built before shipping, decisions were documented, owners were clear, "
            "and cross-functional partners followed through — without adding bureaucracy",
            config, bold_prefix="Program & Project Discipline")
        add_bullet(doc,
            " — acted as business-case clearinghouse: when department heads had half-baked ideas, pressure-tested "
            "them and turned them into board-ready proposals with revenue projections, cost structures, "
            "break-even analysis, and risk assessment that let leadership make fast decisions",
            config, bold_prefix="Business Case Development")
        add_bullet(doc,
            " — built business models for new initiatives: revenue projections, cost structures, break-even "
            "and sensitivity analysis for every new product, location, and investment — napkin math to board-ready",
            config, bold_prefix="Financial Modeling & Analysis")
        add_bullet(doc,
            " — managed full P&L ownership for $4M ARR organization: budget planning, variance analysis, "
            "resource allocation, multi-scenario forecasting across 12 departments and 32 locations",
            config, bold_prefix="P&L Management & Forecasting")
        add_bullet(doc,
            " — led AI-driven workflow transformation: automated billing, scheduling, and reporting processes — "
            "reduced administrative overhead by 40%+, used AI tools daily to accelerate decision-making",
            config, bold_prefix="AI-Augmented Operations")
    else:
        add_body(doc,
            "Directed end-to-end operations for a multi-site healthcare group. Served as the primary "
            "bridge between executive leadership and all operational teams.", config, italic=True, size=Pt(9.5))
        add_bullet(doc,
            " — 8 concurrent due diligence workstreams, integration playbook, Day 1/100 milestones, "
            "100% key talent retention through $17M transition",
            config, bold_prefix="Full-Cycle Acquisition Execution")

    # Earlier Career (add for DoorDash)
    if True:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("Earlier Career")
        run.font.name = config["font"]
        run.font.size = config["size"]
        run.bold = True
        add_bullet(doc, "Digital Strategy Manager (2016–2018) — led digital strategy, campaign analytics, and ROI measurement; built reporting dashboards, optimized $500K+ annual ad spend", config)
        add_bullet(doc, "Client Services Representative (2014–2016) — managed client escalations for enterprise accounts; developed response protocols that reduced resolution time by 30%", config)

    # Education
    add_section_header(doc, "Education", config)
    if company == "DoorDash_Canada":
        add_body(doc, "MBA, Strategy & Finance  |  BSc, Information Technology  |  Post-Bacc Diploma, KPU", config, size=Pt(9), space_after=0)
    else:
        add_body(doc, "Master of Business Administration (MBA) — Strategy, Finance & Operations", config, bold=True, space_after=0)
        add_body(doc, "Post-Baccalaureate Diploma in Business Management — KPU, Surrey, BC", config, size=Pt(9), space_after=0)
        add_body(doc, "Bachelor of Science, Information Technology", config, size=Pt(9), space_after=0)

    # Technical Proficiency
    add_section_header(doc, "Technical Proficiency", config)
    if company == "Methanex":
        add_body(doc,
            "Financial Modelling & Analysis  |  Microsoft Excel (Advanced)  |  Google Workspace  |  "
            "OKR Frameworks  |  ERP Systems  |  CRM Platforms  |  Jira / Confluence  |  Data Visualization",
            config, size=Pt(9))
    elif company == "Hiive":
        add_body(doc,
            "Systems Architecture & Automation  |  EHR / Practice Management Platforms  |  "
            "Google Workspace  |  Data Visualization & KPI Dashboards  |  Financial Modeling  |  "
            "OKR Frameworks  |  CRM Platforms  |  Jira / Confluence  |  AI-Augmented Workflows",
            config, size=Pt(9))
    elif company == "Providence_Healthcare":
        add_body(doc,
            "EHR / Practice Management Platforms  |  Financial Modeling & Budgeting  |  "
            "Data Visualization & KPI Dashboards  |  Google Workspace  |  OKR Frameworks  |  "
            "Project Management Tools  |  Regulatory Compliance Systems  |  Quality Improvement Frameworks",
            config, size=Pt(9))
    elif company == "DoorDash_Canada":
        add_body(doc,
            "ERP/Financial Systems  |  SQL  |  Excel/Google Sheets (advanced modeling)  |  "
            "Data Analysis  |  Project Management Tools",
            config, size=Pt(9))
    elif company == "Practice_Better":
        add_body(doc,
            "Revenue Operations Platforms  |  EHR / Practice Management (Athenahealth, eClinicalWorks)  |  "
            "G Suite  |  CRM Platforms  |  KPI Dashboard Design  |  Financial Modeling  |  "
            "OKR Frameworks  |  Jira / Confluence  |  Data Visualization",
            config, size=Pt(9))
    elif company == "BWZ":
        add_body(doc,
            "Business Modeling (Excel/Sheets)  |  SQL & Data Analytics  |  KPI Dashboard Design  |  "
            "AI-Augmented Workflows  |  Project Management Tools  |  OKR Frameworks  |  "
            "ERP / Financial Systems  |  Data Visualization  |  Cross-Functional Collaboration Tools",
            config, size=Pt(9))
    else:
        add_body(doc,
            "Athenahealth  |  eClinicalWorks  |  CRM Platforms  |  Google Workspace  |  "
            "Financial Modeling  |  OKR Frameworks  |  Jira / Confluence",
            config, size=Pt(9))

    if company == "Methanex":
        role_str = "Director_Strategy"
    elif company == "Hiive":
        role_str = "Associate_Operations_Strategy"
    elif company == "Providence_Healthcare":
        role_str = "Director_Clinical_Operations"
    elif company == "DoorDash_Canada":
        role_str = "Manager_SO_Dasher_Logistics"
    elif company == "UBC":
        role_str = "SrMgr_Strategic_Initiatives"
    elif company == "Practice_Better":
        role_str = "Director_Revenue_Operations"
    elif company == "BWZ":
        role_str = "Strategy_Ops_Manager"
    else:
        role_str = "SrMgr_Integration"

    respath = os.path.join(folder, f"Aman_Kumar_{company}_{role_str}.docx")
    try:
        doc.save(respath)
    except PermissionError:
        respath = os.path.join(lfolder, os.path.basename(respath))
        doc.save(respath)
    doc.save(os.path.join(lfolder, os.path.basename(respath)))
    print(f"Resume: {respath}")

    # --- COVER LETTER ---
    doc = Document()
    set_margins(doc, config["margins"])

    add_body(doc, NAME, config, bold=True, space_after=0)
    cover_contact = doc.add_paragraph()
    cover_contact.paragraph_format.space_after = Pt(0)
    cover_contact.paragraph_format.space_before = Pt(0)
    add_plain_run(cover_contact, f"{PHONE}  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(cover_contact, EMAIL, f"mailto:{EMAIL}", config["font"], Pt(9))
    add_plain_run(cover_contact, "  |  ", config["font"], Pt(9), color="505050")
    add_hyperlink_contact(cover_contact, "LinkedIn", f"https://{LINKEDIN}", config["font"], Pt(9))
    add_body(doc, LOCATION, config, size=Pt(9), space_after=4)
    DATE_LABELS = {"Indeed": "June 20, 2026", "Methanex": "June 21, 2026", "Deloitte": "June 19, 2026", "Hiive": "June 22, 2026", "Providence_Healthcare": "June 22, 2026", "DoorDash_Canada": "June 24, 2026", "UBC": "June 25, 2026", "BWZ": "June 25, 2026"}
    add_body(doc, DATE_LABELS.get(company, "June 22, 2026"), config, space_after=8)

    if company == "Methanex":
        add_body(doc, "Methanex Corporation", config, space_after=0)
        add_body(doc, "1800 Waterfront Centre, 200 Burrard Street", config, space_after=0)
        add_body(doc, "Vancouver, BC V6C 3M1", config, space_after=8)
        add_body(doc, "Re: Director, Strategy", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "Methanex is the world's largest methanol producer, navigating volatile commodity markets, "
            "the integration of a $2.05B acquisition, and a strategic shift toward low-carbon solutions. "
            "The Director, Strategy role sits at the center of these dynamics. This is exactly where I deliver the most value.\n\n"
            "I spent eight years building the strategic infrastructure of a multi-site organization from scratch. "
            "When I joined, there were 3 people and no playbook. When I left, it was 70 people across 32 locations "
            "with a $17M acquisition that I directed end-to-end. The strategic planning system I designed — annual "
            "strategy sessions cascading through quarterly OKRs with board-level reporting — gave executive leadership "
            "real-time visibility into execution across every location. The financial models I built governed resource "
            "allocation across 12 departments. The acquisition I directed involved 8 workstreams, diligence across "
            "finance, legal, and operations, and an integration that retained 100% of our key talent.\n\n"
            "What made this possible was not a pre-existing framework. I built it — the strategy cycle, the governance "
            "rhythms, the valuation models, the integration playbook — all from zero. I learned that strategy is not "
            "a document you produce once a year. It is a living process that must connect the Board room to the front "
            "line, and it requires someone who can operate at both altitudes without losing coherence.\n\n"
            "I understand Methanex operates in a different industry. But strategy is fractal — the core problems are "
            "the same: Where do we allocate capital? How do we grow? What risks do we manage? How do we align a global "
            "organization around a shared plan? I have solved these problems in my domain, and I am ready to solve them in yours.\n\n"
            "I would welcome the opportunity to discuss how my experience building strategic and financial infrastructure "
            "can support Methanex's next phase of global leadership."
        )
    elif company == "Hiive":
        add_body(doc, "Hiive", config, space_after=0)
        add_body(doc, "Vancouver, BC (HQ)", config, space_after=8)
        add_body(doc, "Re: Associate, Operations Strategy", config, bold=True, space_after=8)

    elif company == "Providence_Healthcare":
        add_body(doc, "Providence Health Care", config, space_after=0)
        add_body(doc, "Burnaby, BC", config, space_after=8)
        add_body(doc, "Re: Director, Clinical and Operations (LTC PHC & FHA)", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Committee,\n\n"
            "I spent 8 years leading multi-site healthcare operations — and the patient population I served? "
            "Seniors. I have done exactly what this role requires.\n\n"
            "At SkyflyMD, I served as the de facto Director of Operations for a multi-site healthcare organization "
            "that grew from 3 to 70 people across 32 locations. The practice was geriatrics-focused — I managed operations "
            "alongside a geriatrician, coordinating care across clinic locations and senior care homes, managing "
            "billing and compliance for aging populations, and building the systems that supported quality care "
            "for older adults. I managed the full P&L ($4M ARR). I built the operational infrastructure from "
            "scratch — scheduling, billing, compliance, quality assurance, reporting systems. I led interdisciplinary "
            "teams across 12 departments. I directed the $17M acquisition that integrated 8 separate operational "
            "systems into one without losing a single key team member.\n\n"
            "What drew me to Providence is the alignment between my experience and your mission. Multi-site seniors care "
            "at Chenchenstway and your PHC/FHA long-term care sites faces exactly the operational challenges I have been "
            "solving for the past 8 years: how to scale quality care across multiple locations serving aging populations "
            "without losing the personalized, compassionate approach that defines your organization.\n\n"
            "I hold an MBA with 8 years of progressive healthcare operations leadership — including direct experience "
            "serving geriatric populations in multi-facility settings. My combination of formal business education and "
            "hands-on seniors care operations experience provides the equivalent foundation this role requires.\n\n"
            "I would welcome the opportunity to discuss how my experience building and leading multi-site healthcare "
            "operations for aging populations can support Providence and Fraser Health's vision for seniors care."
        )
    elif company == "DoorDash_Canada":
        add_body(doc, "DoorDash Canada", config, space_after=0)
        add_body(doc, "Vancouver, BC / Toronto, ON", config, space_after=8)
        add_body(doc, "Re: Manager, Strategy & Operations – Dasher & Logistics Canada", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Team,\n\n"
            "I built a business from nothing — scaled it to 70 people, 32 locations, "
            "managed every dollar of the P&L, and delivered a $17M exit. Along the way, "
            "I learned something that applies directly to this role: how you design pay "
            "and incentive systems determines whether your frontline workforce treats "
            "their work as a transaction or a partnership.\n\n"
            "At my company, I owned the complete compensation architecture for a "
            "distributed workforce across 32 locations. Every market had different "
            "dynamics — different cost of labour, different competitive pressure, "
            "different worker expectations. I built pay structures that balanced worker "
            "earnings with business cost efficiency, and I learned the behavioral "
            "economics of what actually motivates people in gig-like environments. "
            "Small changes in incentive design drove outsized shifts in performance.\n\n"
            "That experience maps directly to the challenge you're solving on the "
            "Dasher & Logistics team. DoorDash's marketplace depends on getting the "
            "pay and incentive equation right — attracting and retaining high-quality "
            "Dashers, at efficient spend, across diverse regional markets. It's a "
            "marketplace optimization problem that I've lived at operational scale.\n\n"
            "I'm drawn to DoorDash because your values match how I actually operate. "
            "\"Be an Owner\" — that's how I ran my business. \"Operate at the lowest "
            "level of detail\" — I still know the weekly cost per visit at every "
            "location I built. \"Bias for Action\" — I don't wait for perfect data "
            "to make a decision. I launch, test, iterate, and compound the wins.\n\n"
            "I'm ready to bring that owner-operator mindset to the Dasher & Logistics "
            "team. I understand the mechanics of labour supply, the tension between "
            "cost and quality, and the power of well-designed incentives.\n\n"
            "Let's talk."
        )
    elif company == "Hiive":
        add_body(doc, "Hiive", config, space_after=0)
        add_body(doc, "Vancouver, BC (HQ)", config, space_after=8)
        add_body(doc, "Re: Associate, Operations Strategy", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "I spent eight years doing exactly what this role describes: embedding with a revenue-generating team, "
            "identifying bottlenecks in the operational lifecycle, and deploying automated solutions to clear them.\n\n"
            "When I joined SkyflyMD, there was no operational infrastructure — just a team operating on willpower. "
            "By the time I directed the $17M exit, we had:\n\n"
            "- A complete tech stack (EHR, billing, scheduling, analytics) that I designed and implemented from scratch\n"
            "- Real-time KPI dashboards replacing manual spreadsheets\n"
            "- Automated workflows that reduced administrative overhead by 40%\n"
            "- A scalable operating system that supported 70 people across 32 locations — without adding operational complexity\n\n"
            "I did not inherit these systems. I built them. That is what systems builder means to me.\n\n"
            "Hiive is doing the same thing for the private market that I did for healthcare operations: replacing opacity "
            "with transparency, manual processes with automation, fragmentation with integration. The private secondary "
            "market has been running on brokers and spreadsheets for too long. You are building the infrastructure "
            "that changes that.\n\n"
            "I want to help you clear the bottlenecks.\n\n"
            "I am based in Vancouver and ready to be in your HQ five days a week."
        )
    elif company == "UBC":
        add_body(doc, "UBC Human Resources", config, space_after=0)
        add_body(doc, "University of British Columbia, Vancouver, BC", config, space_after=8)
        add_body(doc, "Re: Senior Manager, Strategic Initiatives and Engagement", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Committee,\n\n"
            "UBC's Shaping UBC's Next Century plan articulates an ambitious mandate: deliver strategic priorities "
            "through disciplined resource allocation, stakeholder engagement, and institutional effectiveness. "
            "The Senior Manager, Strategic Initiatives and Engagement role sits at the center of this mandate. "
            "This is exactly where I deliver the most value.\n\n"
            "I spent eight years building the strategic and operational infrastructure for a multi-site organization "
            "from scratch. When I joined, there were 3 people, no governance structure, and no strategic planning "
            "process. When I left, it was 70 people across 32 locations with a $17M acquisition that I directed "
            "end-to-end. The strategic planning system I built --- annual strategy sessions cascading through "
            "departmental OKRs, board-level reporting, and quarterly performance reviews --- gave executive leadership "
            "real-time visibility into execution across every location. The budget frameworks I designed governed "
            "resource allocation across 12 departments. The acquisition I led involved 8 concurrent workstreams and "
            "an integration that retained 100% of key talent.\n\n"
            "What made this possible was not a pre-existing playbook. I created it --- the governance rhythms, "
            "the strategic planning cycles, the financial models, the stakeholder engagement frameworks, the policy "
            "documentation systems --- all from scratch, across a complex, multi-jurisdictional operating environment. "
            "I learned that strategic leadership in complex institutions is not about having all the answers. It is "
            "about building the process that surfaces the right questions, engages the right stakeholders, and delivers "
            "the right outcomes.\n\n"
            "UBC faces exactly the challenges I have solved: how to allocate constrained resources across competing "
            "strategic priorities, how to align diverse stakeholders around shared goals, how to build governance "
            "frameworks that satisfy both institutional accountability and external transparency, how to translate "
            "strategic intent into operational reality across a large, complex organization. I have navigated this "
            "complexity in healthcare. I am ready to navigate it at UBC.\n\n"
            "I am drawn to UBC because it is the intellectual backbone of British Columbia's future. I want to "
            "contribute to that future as someone who builds --- not as an observer, but as a practitioner who has "
            "delivered strategic infrastructure at scale. My MBA, my 8 years of progressive operational leadership, "
            "and my $17M acquisition experience have prepared me to do exactly what this role requires.\n\n"
            "I would welcome the opportunity to discuss how my experience building strategic and operational "
            "infrastructure can support UBC's next chapter of academic excellence and institutional impact."
        )
    elif company == "Practice_Better":
        add_body(doc, "Practice Better", config, space_after=0)
        add_body(doc, "Toronto, ON (Remote Canada)", config, space_after=8)
        add_body(doc, "Re: Director of Revenue Operations", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "I build revenue operations infrastructure. I've done it from zero, scaled it to $4M ARR, "
            "and delivered a $17M exit. That's the exact journey Practice Better is on right now --- "
            "and I want to be the operator who builds what comes next.\n\n"
            "When I joined SkyflyMD, there was no revenue operations system. No billing infrastructure, "
            "no scheduling automation, no KPI visibility. By the time I directed the $17M exit, we had "
            "a complete revenue operations stack spanning 32 locations: EHR, billing, RCM, analytics, "
            "and forecasting --- all built in-house, all scaled from zero to support 70 employees and "
            "$4M ARR. I didn't inherit a system. I built one.\n\n"
            "What makes this relevant to Practice Better is what I learned building that system: in healthcare "
            "SaaS, revenue operations is not just about pipeline management. It's about aligning clinical "
            "workflows with financial outcomes --- making sure the GTM engine and the product engine speak "
            "the same language. That's the gap I've seen most scaling healthcare SaaS companies struggle with, "
            "and that's exactly what I've solved.\n\n"
            "At SkyflyMD, I also managed the full P&L --- $4M ARR across 12 departments. I built the "
            "forecasting models, the variance analysis frameworks, and the board-level reporting that gave "
            "executive leadership real-time visibility. I know what it means to own the numbers and be "
            "accountable for them.\n\n"
            "I'm targeting Director of Revenue Operations roles at scaling healthcare SaaS companies because "
            "this is the exact problem I've already solved. Practice Better is at the stage where the founder "
            "can't own revenue ops alone anymore --- and that's where I come in. I build the function, hire "
            "the team, and create the infrastructure that makes the next phase of growth inevitable.\n\n"
            "I would welcome the opportunity to discuss how my experience building and scaling revenue "
            "operations for healthcare SaaS can support Practice Better's next chapter."
        )
    elif company == "BWZ":
        add_body(doc, "Black & White Zebra", config, space_after=0)
        add_body(doc, "Vancouver, BC (Remote)", config, space_after=8)
        add_body(doc, "Re: Strategy & Operations Manager", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "I read your JD and felt like someone had transcribed my career. 'Own the end-to-end build of zero-to-one initiatives.' "
            "'Bring order to the messy middle.' 'Impatient with meetings that don't move toward a decision.' "
            "That's not a job description — that's how I've operated for the last 8 years.\n\n"
            "I built a multi-site organization from 3 people to 70, across 32 locations, from zero to $4M ARR. "
            "I led the technology transformation, managed every dollar of the P&L, and directed a $17M exit. "
            "Along the way, I became the person people came to when they had a half-baked idea that needed "
            "pressure-testing — building the business case, running the costing, designing the go-to-market, "
            "and delivering the board-ready package that let leadership say yes or no with confidence.\n\n"
            "That's exactly what BWZ needs right now. You're at that sweet spot — 70 people, diversifying revenue, "
            "expanding beyond content into SaaS, and you need someone who can install project discipline "
            "without killing the startup soul. I've already walked this path. I know which processes matter "
            "and which ones just add bureaucracy.\n\n"
            "What draws me to BWZ specifically: you're bootstrapped and independent. You're not chasing unicorn "
            "valuations — you're building something sustainable. That's exactly how I built my company. "
            "And your CEO Ben Aston started as a solo blogger and built this from nothing — "
            "that's a founder story I deeply understand because I lived my own version of it.\n\n"
            "I'd welcome the chance to talk about how I can help BWZ bring order to the messy middle "
            "while keeping the builder spirit that got you here."
        )
    else:
        add_body(doc, "Indeed", config, space_after=0)
        add_body(doc, "Vancouver, BC", config, space_after=8)
        add_body(doc, "Re: Sr. Manager, Integration & Business Acceleration — Reference ID: 47053", config, bold=True, space_after=8)

        body = (
            "Dear Hiring Manager,\n\n"
            "Your mission is to help people get jobs. Mine is to build the operational infrastructure "
            "that makes organizations scalable, efficient, and ready for their next phase of growth.\n\n"
            "I joined SkyflyMD when it was 3 people in a single location. When I left, it was 70 people "
            "across 32 locations, operating with systems I designed from scratch. The defining moment of "
            "that journey was directing our $17M acquisition — structuring 8 concurrent due diligence "
            "workstreams across finance, legal, operations, and provider contracts. I built the integration "
            "playbook, tracked Day 1 readiness and Day 100 milestones, consolidated 8 separate operational "
            "systems into one unified platform, and retained 100% of our key talent through the transition.\n\n"
            "What made that possible was not a pre-existing framework. There was no playbook. I created it — "
            "the governance rhythms, the reporting cadences, the escalation paths, the decision-making "
            "frameworks — all from scratch, across 5 clinic groups operating under different state regulations.\n\n"
            "What drew me to Indeed is that you are data-driven without being jargon-heavy, and mission-focused "
            "without being sentimental. That is how I operate. Your 'job seeker first' value resonates because "
            "it mirrors how I have always made operational decisions — starting with the end-user and working backward.\n\n"
            "I am not looking for a role that requires a playbook to exist before I start. I am looking for "
            "one that needs a playbook written.\n\n"
            "I would welcome the opportunity to discuss how my experience building and integrating multi-site "
            "operations can support Indeed's continued growth through M&A. Thank you for your time and consideration."
        )

    add_body(doc, body, config, space_after=0)
    add_body(doc, "", config, space_after=0)
    add_signature(doc, config)

    clpath = os.path.join(folder, f"Cover_Letter_{company}_{role_str}.docx")
    try:
        doc.save(clpath)
    except PermissionError:
        clpath = os.path.join(lfolder, os.path.basename(clpath))
        doc.save(clpath)
    doc.save(os.path.join(lfolder, os.path.basename(clpath)))
    print(f"Cover:   {clpath}")

if __name__ == "__main__":
    company = sys.argv[1] if len(sys.argv) > 1 else "Methanex"
    generate(company)
