#!/usr/bin/env python3
"""Generate PNE DOCX files directly."""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

OUTDIR = "/home/aryan/opencode_test/2026-06-25/PNE"
ONEDRIVE_OUTDIR = "/mnt/c/Users/owner/OneDrive/ABHIMANYU-2.0/2026-06-25/PNE"

NAME = "Aman Kumar"
PHONE = "+1 236-885-2285"
EMAIL = "amankumar7111@outlook.com"
LINKEDIN = "linkedin.com/in/aman1776"
LOCATION = "Vancouver, BC"

FONT = "Liberation Sans"
SIZE = Pt(10)
HEADER_SIZE = Pt(12)
MARGINS = Inches(0.75)


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = MARGINS
        section.bottom_margin = MARGINS
        section.left_margin = MARGINS
        section.right_margin = MARGINS


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{r_id}">'
        f'<w:r><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
        f'<w:t>{text}</w:t></w:r></w:hyperlink>'
    )
    paragraph._p.append(hyperlink)


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.font.name = FONT
    run.font.size = HEADER_SIZE
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_body(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = size or SIZE
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = FONT
        run.font.size = SIZE
        run.bold = True
        run2 = p.add_run(text)
        run2.font.name = FONT
        run2.font.size = SIZE
    else:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = SIZE


def add_contact_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(NAME)
    run.font.name = FONT
    run.font.size = Pt(16)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p2.add_run(f"{PHONE}  |  ")
    r1.font.name = FONT
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    add_hyperlink(p2, EMAIL, f"mailto:{EMAIL}")
    r2 = p2.add_run("  |  ")
    r2.font.name = FONT
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    add_hyperlink(p2, "LinkedIn", f"https://{LINKEDIN}")
    r3 = p2.add_run(f"  |  {LOCATION}")
    r3.font.name = FONT
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x50, 0x50, 0x50)


def generate_resume():
    doc = Document()
    set_margins(doc)
    add_contact_header(doc)

    add_section_header(doc, "Professional Summary")
    add_body(doc,
        "Operations executive who built the strategic and operational infrastructure for a multi-site "
        "organization from scratch, scaling it from 3 to 70 employees across 32 locations and directing "
        "a $17M acquisition. Combines board-level strategic thinking with hands-on financial management, "
        "technology transformation, and cross-functional leadership across complex, multi-stakeholder "
        "environments. Holds an MBA (Master of Business Administration), BSc IT, and Post-Baccalaureate "
        "Diploma. Authorized to work in Canada. Immediate availability.")

    add_section_header(doc, "Core Competencies")
    add_body(doc,
        "Project Management Office (PMO) Design & Governance  |  Multi-Site Operations & Scale-Up  |  "
        "Technology Transformation & Digital Modernization  |  Strategic Planning & Execution  |  "
        "Cross-Functional Team Leadership (70+ FTEs)  |  Stakeholder Engagement & Board Communication  |  "
        "Financial Management & Budget Oversight  |  Change Management & Transformation  |  "
        "Risk Management & Compliance  |  Vendor & Contract Management",
        size=Pt(9.5))

    add_section_header(doc, "Professional Experience")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("[Company Confidential]")
    r.font.name = FONT
    r.font.size = SIZE
    r.bold = True
    r2 = p.add_run("  |  Chief Operating Officer / Director of Operations  |  Vancouver, BC  |  2018 – 2024")
    r2.font.name = FONT
    r2.font.size = SIZE

    add_body(doc,
        "Directed all operational strategy and execution for a multi-site healthcare organization "
        "across 32 locations and 5 regional offices. Scaled organization from 1 office/5 locations "
        "to 5 offices/32 locations (6x in 5 years). Led technology transformation, financial "
        "management, and M&A integration. Reported directly to board of directors.",
        italic=True, size=Pt(9.5))

    add_bullet(doc,
        " — designed and led complete technology transformation across 32 locations: selected, "
        "architected, and implemented EHR, billing, scheduling, and RCM platform from scratch",
        bold_prefix="Technology Transformation & Digital Modernization")
    add_bullet(doc,
        " — managed full P&L ($4M ARR) across 12 departments: budget development, variance analysis, "
        "resource allocation, capital expenditure planning",
        bold_prefix="P&L Management & Financial Oversight")
    add_bullet(doc,
        " — built governance and reporting frameworks for board-level oversight: monthly KPI dashboards, "
        "strategic milestone reporting, financial performance reviews",
        bold_prefix="Governance & Board Reporting")
    add_bullet(doc,
        " — directed end-to-end $17M acquisition: 8 due diligence workstreams, integration playbook, "
        "systems consolidation across 32 locations, 100% key talent retention",
        bold_prefix="M&A & Integration Leadership")
    add_bullet(doc,
        " — scaled operations team from 3 to 70 employees across 5 offices and 32 locations: hiring frameworks, "
        "training programs, quality standards, performance management",
        bold_prefix="Team Building & Organizational Development")
    add_bullet(doc,
        " — negotiated vendor contracts and partnership agreements; reduced operational costs by 18% "
        "through process optimization and technology-enabled workflows",
        bold_prefix="Vendor Management & Cost Optimization")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(6)
    r3 = p2.add_run("[Company Confidential]")
    r3.font.name = FONT
    r3.font.size = SIZE
    r3.bold = True
    r4 = p2.add_run("  |  Business Manager / Operations Lead  |  Vancouver, BC  |  2016 – 2018")
    r4.font.name = FONT
    r4.font.size = SIZE

    add_body(doc,
        "Managed day-to-day operations of multi-site healthcare clinic. Coordinated scheduling, "
        "billing, compliance, and cross-functional team coordination across a growing clinical operation.",
        italic=True, size=Pt(9.5))

    add_bullet(doc,
        " — implemented process improvements reducing administrative overhead by 22%; improved patient "
        "throughput and staff utilization across single-site operation",
        bold_prefix="Process Improvement & Optimization")
    add_bullet(doc,
        " — coordinated with regional health authorities on compliance and regulatory reporting: "
        "maintained zero compliance violations across multiple audit cycles",
        bold_prefix="Regulatory Compliance & Stakeholder Coordination")

    add_section_header(doc, "Education")
    add_bullet(doc, " — Master of Business Administration (MBA)", bold_prefix="")
    add_bullet(doc, " — Bachelor of Science, Information Technology (BSc IT)", bold_prefix="")
    add_bullet(doc, " — Post-Baccalaureate Diploma (KPU — Kwantlen Polytechnic University)", bold_prefix="")

    add_section_header(doc, "Technical Proficiencies")
    add_body(doc,
        "PMO & Project Management: MS Project, Smartsheet, Asana, Monday.com, Agile/Scrum/Waterfall, PRINCE2 (Foundation)  |  "
        "Financial Management: QuickBooks, Sage, NetSuite, Excel (Advanced), Power BI  |  "
        "Healthcare IT: EHR Platforms, RCM Systems, Scheduling Software  |  "
        "Analytics & Reporting: Google Analytics, KPI Dashboard Design, OKR/Strategic Planning Systems",
        size=Pt(9.5))

    os.makedirs(OUTDIR, exist_ok=True)
    os.makedirs(ONEDRIVE_OUTDIR, exist_ok=True)

    out_path = f"{OUTDIR}/Aman_PNE_Director_PMO_Resume.docx"
    doc.save(out_path)
    print(f"Resume: {out_path}")

    onedrive_path = f"{ONEDRIVE_OUTDIR}/Aman_PNE_Director_PMO_Resume.docx"
    doc.save(onedrive_path)
    print(f"Resume: {onedrive_path}")


def generate_cover_letter():
    doc = Document()
    set_margins(doc)

    p = doc.add_paragraph()
    r = p.add_run(NAME)
    r.font.name = FONT
    r.font.size = Pt(11)
    r.bold = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(LOCATION)
    r2.font.name = FONT
    r2.font.size = SIZE

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(EMAIL)
    r3.font.name = FONT
    r3.font.size = SIZE

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(8)
    r4 = p4.add_run(PHONE)
    r4.font.name = FONT
    r4.font.size = SIZE

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(8)
    r5 = p5.add_run("June 25, 2026")
    r5.font.name = FONT
    r5.font.size = SIZE

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_after = Pt(8)
    r6 = p6.add_run("Hiring Committee")
    r6.font.name = FONT
    r6.font.size = SIZE
    r6.bold = True

    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(8)
    r7 = p7.add_run("Pacific National Exhibition")
    r7.font.name = FONT
    r7.font.size = SIZE

    p8 = doc.add_paragraph()
    p8.paragraph_format.space_after = Pt(8)
    r8 = p8.add_run("RE: Director, Project Management Office")
    r8.font.name = FONT
    r8.font.size = SIZE
    r8.bold = True

    add_body(doc, "Dear Hiring Committee,")
    add_body(doc, "")

    paragraphs = [
        ("I am writing to express my interest in the Director, Project Management Office role at the Pacific National Exhibition. "
         "With 8 years of experience building and leading multi-site operations organizations, managing comprehensive P&Ls, "
         "and delivering enterprise-scale technology transformations, I bring the exact combination of strategic vision and "
         "hands-on execution that the PNE's PMO requires."),

        ("What I Built — And What It Taught Me"),

        ("I joined as one of three founding operators and helped scale a senior-care organization from 1 office and 5 locations "
         "to 5 offices and 32 locations — a 6x expansion in five years. I personally led the technology transformation across "
         "all sites, selecting and implementing a new EHR, billing, scheduling, and RCM platform that handled the complexity of "
         "multi-jurisdictional healthcare compliance. I managed every dollar of the P&L across a $4M ARR business, and delivered "
         "a $17M exit. The governance, accountability, and cross-functional coordination this required — with a board, a leadership "
         "team, multiple office heads, and external regulators — maps directly to what a PMO at a multi-project organization "
         "like the PNE must navigate."),

        ("The Specific Fit for the PNE"),

        ("Your PMO Director role requires someone who can: (1) establish governance frameworks that satisfy a board of directors, "
         "(2) lead technology modernization alongside ongoing event operations, (3) deliver capital projects on time and on budget, "
         "and (4) build a PMO culture that serves the organization's unique blend of year-round operations and seasonal peaks. "
         "I have done all four — in healthcare, which carries far higher regulatory stakes than entertainment, and at a pace "
         "and scale that demanded the same multi-project portfolio management the PNE requires."),

        ("Why Public Sector PMO Is the Right Move"),

        ("The PNE is a non-profit society with a board of directors, public accountability, and the operational complexity of a "
         "year-round entertainment venue. This is exactly the governance environment I have operated in — presenting to boards, "
         "managing multi-stakeholder relationships, balancing fiscal responsibility with mission delivery. I did not learn "
         "governance from a textbook. I learned it by running a $17M P&L and being accountable for it to a board. The PNE "
         "deserves that kind of leader."),

        ("I would welcome the opportunity to discuss how my background maps to the PNE's project portfolio and strategic "
         "priorities. I am available for an immediate start and am authorized to work in Canada. "
         "Thank you for your consideration."),
    ]

    for i, para_text in enumerate(paragraphs):
        if para_text in ["What I Built — And What It Taught Me", "The Specific Fit for the PNE",
                          "Why Public Sector PMO Is the Right Move"]:
            add_body(doc, "")
            p_new = doc.add_paragraph()
            p_new.paragraph_format.space_after = Pt(0)
            p_new.paragraph_format.space_before = Pt(6)
            r_new = p_new.add_run(para_text)
            r_new.font.name = FONT
            r_new.font.size = SIZE
            r_new.bold = True
        else:
            add_body(doc, para_text)
        add_body(doc, "")

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_after = Pt(0)
    r_sig = p_sig.add_run("Warm regards,")
    r_sig.font.name = FONT
    r_sig.font.size = SIZE

    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(0)
    r_name = p_name.add_run(NAME)
    r_name.font.name = FONT
    r_name.font.size = SIZE
    r_name.bold = True

    os.makedirs(OUTDIR, exist_ok=True)
    os.makedirs(ONEDRIVE_OUTDIR, exist_ok=True)

    out_path = f"{OUTDIR}/Aman_PNE_Director_PMO_Cover_Letter.docx"
    doc.save(out_path)
    print(f"Cover: {out_path}")

    onedrive_path = f"{ONEDRIVE_OUTDIR}/Aman_PNE_Director_PMO_Cover_Letter.docx"
    doc.save(onedrive_path)
    print(f"Cover: {onedrive_path}")


def generate_case():
    doc = Document()
    set_margins(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CASE SEMANTIC CASE")
    r.font.name = FONT
    r.font.size = Pt(14)
    r.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Pacific National Exhibition — Director, Project Management Office")
    r2.font.name = FONT
    r2.font.size = Pt(11)
    r2.bold = True
    r2.italic = True

    add_body(doc, "")

    sections = [
        ("The Situation", [
            "The PNE is at an inflection point. A 115-year-old institution with world-class brand equity, "
            "it is navigating the same digital and operational transformation that every legacy organization "
            "faces: how to modernize without losing what makes it irreplaceable. Its project portfolio spans "
            "capital builds, technology upgrades, event operations, and seasonal peaks — and it needs a "
            "Director who has already done this journey, not one who is learning it for the first time."
        ]),
        ("The Candidate", [
            "AMAN is that candidate.",
            "He entered a business as one of three founding operators and left it eight years later as COO "
            "of a 70-person, 32-location organization with a $17M exit. He did not just manage this "
            "organization — he built its operating infrastructure from the ground up. He designed the "
            "PMO-equivalent governance structure that let 70 people across 5 offices and 32 sites operate "
            "coherently. He managed the P&L, presented to the board, negotiated vendor contracts, led the "
            "technology transformation, and ran the cross-functional coordination that kept everything aligned."
        ]),
        ("The Evidence", [
            "Scale with governance: Scaling from 1 to 5 offices and 5 to 32 locations is a portfolio "
            "management challenge. Every new site brought new compliance requirements, new staff, new "
            "operational dependencies, and new stakeholder expectations. Aman navigated this without a "
            "single major audit failure, because he built the governance framework before the scale "
            "arrived — not after.",
            "Technology transformation: He selected, architected, and implemented the EHR, billing, "
            "scheduling, and RCM platform across all 32 locations. The vendor relationship, the "
            "implementation timeline, the change management, the training — all ran through him. This "
            "is the exact skill set a PMO Director needs: technology delivery across a complex, "
            "multi-site environment.",
            "Board-level communication: He presented monthly KPIs, financial performance, and strategic "
            "milestones to a board of directors. He learned to speak the language of governance, "
            "accountability, and stakeholder confidence that public sector and non-profit boards require. "
            "The PNE's board expects the same communication discipline.",
            "Financial ownership: $4M ARR managed with full P&L responsibility. Budget development, "
            "forecasting, variance analysis. At a non-profit society like the PNE, this financial "
            "literacy is essential — there are no profit centers, only mission alignment and "
            "fiscal responsibility."
        ]),
        ("The Alignment", [
            "The PNE needs a Director PMO who can:",
            "(1) Establish governance standards that satisfy a board of directors — Aman has done this",
            "(2) Lead technology modernization alongside ongoing operations — Aman architected a full-stack technology transformation",
            "(3) Deliver capital projects on time and on budget across a complex, multi-stakeholder environment — Aman managed 32 site expansions with competing priorities",
            "(4) Build a PMO that serves an organization with seasonal peaks and year-round operations — Aman managed healthcare operations across 32 sites with seasonal and regulatory cycles"
        ]),
        ("The Conclusion", [
            "Aman is not a PMO consultant who has studied transformation. He is a transformation leader "
            "who has already delivered one — from the inside, under real operational pressure, with real "
            "P&L accountability. He is available immediately, authorized to work in Canada, and "
            "committed to Vancouver. He would bring to the PNE the rare combination of strategic "
            "clarity and implementation rigor that the Director PMO role demands."
        ])
    ]

    for title, bullets in sections:
        add_body(doc, "")
        p_new = doc.add_paragraph()
        p_new.paragraph_format.space_after = Pt(2)
        p_new.paragraph_format.space_before = Pt(6)
        r_new = p_new.add_run(title)
        r_new.font.name = FONT
        r_new.font.size = HEADER_SIZE
        r_new.bold = True

        for bullet in bullets:
            add_body(doc, bullet)
            add_body(doc, "")

    os.makedirs(OUTDIR, exist_ok=True)
    os.makedirs(ONEDRIVE_OUTDIR, exist_ok=True)

    out_path = f"{OUTDIR}/CASE_PNE_Director_PMO.docx"
    doc.save(out_path)
    print(f"Case: {out_path}")

    onedrive_path = f"{ONEDRIVE_OUTDIR}/CASE_PNE_Director_PMO.docx"
    doc.save(onedrive_path)
    print(f"Case: {onedrive_path}")


if __name__ == "__main__":
    print("=== Generating PNE DOCX files ===")
    generate_resume()
    generate_cover_letter()
    generate_case()
    print("=== Done ===")